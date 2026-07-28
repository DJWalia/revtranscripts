
import streamlit as st

st.title("Rev Transcript Formatter")
st.write(
    "Input file (Word docs only, multiple files allowed (200 MB Max Total Size)."
)

from docx import Document
from docx.shared import Pt, Inches
import io

def get_docx_bytes(document_object):
    bio = io.BytesIO()
    document_object.save(bio)
    return bio.getvalue()

def convert_transcript(file):

    doc = Document(file)
    new_doc = Document()

    style = new_doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    sections = new_doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    p = new_doc.add_paragraph()

    top_labels = [
        "Platform:",
        "Date:",
        "Event Title:",
        "Length:",
        "Type of Recording:",
        "Transcript:",
        "Saved Clip:",
        "Public Link:"
    ]
    for label in top_labels:
        run_label = p.add_run(label)
        run_label.bold = True
        run_space = p.add_run(" ")
        run_space.bold = False
        run_space.add_break()

    run_video = p.add_run("(VIDEO)")
    run_video.add_break()
    run_video.add_break()

    run_transcript = p.add_run("TRANSCRIPT:")
    run_transcript.bold = True
    run_transcript.add_break()

    p.add_run().add_break()

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    i = 0
    while i < len(paragraphs):
        
        line = paragraphs[i]
        if "(" in line and ")" in line:
            try:
                speaker_part, rest = line.split("(", 1)
                timestamp = rest.split(")")[0].strip()
                speaker_name = speaker_part.strip().upper()

                dialogue = ""
                if i + 1 < len(paragraphs) and "(" not in paragraphs[i + 1]:
                    dialogue = paragraphs[i + 1].strip()
                    i += 1

                run_speaker = p.add_run(f"[{timestamp}] {speaker_name}: ")
                run_speaker.bold = True
                p.add_run(dialogue)
                p.add_run().add_break()
                p.add_run().add_break()

            except Exception:
                p.add_run(line)
                p.add_run().add_break()
                p.add_run().add_break()

        else:
            p.add_run(line)
            p.add_run().add_break()
            p.add_run().add_break()

        i += 1

    input_file_name = str(file.name)
    input_file_name_docx = input_file_name.replace(".docx", "")
    formatted_name = input_file_name_docx + ' formatted.docx'
    formatted_docx = get_docx_bytes(new_doc)
    return formatted_docx, formatted_name

input_file = st.file_uploader("Choose a .docx file", type=['docx'], accept_multiple_files=True)

if input_file is not None:
    for file in input_file:
        formatted_docx, formatted_name = convert_transcript(file)
        label = f"Download Formatted Transcript ({formatted_name})"
        st.download_button(label=label, data=formatted_docx, file_name= formatted_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
